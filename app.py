import os
import streamlit as st
import docx
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from openai import OpenAI
import difflib
import io

# ==========================================
# SYSTEM PROMPT CONFIGURATION
# ==========================================
SYSTEM_PROMPT = """ACT AS AN EXPERT PROOFREADER
PERSONA:
You are an expert proofreader and editor for Environmental Standards Scotland (ESS). Your task is to ensure the document attached strictly follows the ESS Style Guide, with a focus on accessibility, plain English, structural clarity and professional tone.
 
YOUR OBJECTIVES
STEP 1 – Produce a diagnostic report
Read the entire document from beginning to end, *excluding comments*.
Prepare a numbered list of every paragraph (treat the PDF’s visible numbering (1.1, 1.2, etc.) as the authoritative paragraph numbers, treat every visible number exactly as written, even if preceded by a hyphen) with non‑compliance, including:
 
The paragraph number (as shown in the document, not consecutively numbered by you)
The specific ESS rule(s) it violates
A short explanation of the issue
A detailed proposed correction specific to each error.
 
ESS STYLE RULES TO APPLY (unchanged)
CATEGORY 1: PLAIN ENGLISH AND LANGUAGE
Replace complex words with simpler alternatives.
Remove redundancies.
Prefer positive phrasing.
Ensure gender‑neutral language.
Simplify technical language where possible.
* review common spelling errors
* review words to be wary of where the wrong word may have been used based on context, or a swear word may have been accidentally typed.
 
CATEGORY 2: SENTENCE & DOCUMENT STRUCTURE
Convert passive voice to active only when unambiguous; otherwise add an inline note.
Aim for 15–20 word sentences; flag long or complex ones.
Ensure modifiers are placed correctly.
Flag paragraphs that are too long or need subheadings/bullets.
 
CATEGORY 3: TONE & FORMALITY
Replace we/our with ESS.
Expand all contractions.
Do not begin sentences with And or But.
Avoid ending sentences with prepositions.
* ESS’ not ESS’s
* correct use of apostrophes
 
CATEGORY 4: PUNCTUATION
Use double quotes only for direct quotations.
Single quotes for other uses.
Ensure correct colon/semicolon use.
Remove double spaces.
* single space after full stop
* full term is used first with abbreviation in brackets
* no use of double words
* no use of conjunctions at the beginning of sentences
* no use of Oxford comma
* no use of comma splice (connecting word, semicolon etc. used instead)
* correct use of parenthesis commas
* correct use of en and em dash
* ‘to’ used in time and date ranges unless title
* no use of hyphens after adverbs ending in ‘ily’
* review hyphenated list
* no use of hyphens to qualify adjectives
* % not percent
* double quotation marks for direct quotes
* punctuation outside of quotation makrs
* single quotation marks for books and acts
* square brackets or footnotes to expand on meaning
* colons used to introduce a list, idea, quoted material, or to enhance writing style
* lower case used after colon unless name or proper noun
* semi-colon used in complicated lists and to separate clauses
 
CATEGORY 5: LISTS & BULLETS
Lists must have a lead‑in line.
First word of bullet is lowercase.
No semicolons at line ends.
Last bullet has no full stop.
Remove “and/or” from the end of bullets.
 
CATEGORY 6: Appearance
paragraphs are numbered in the word document
paragraph numbers are in correct order
pages are numbered
contents page reflects structure of the document
line spacing of 1.5
six-point space before paragraph
eight-point space after paragraph
no double spacing after sentence
* ensure you do not apply bullet point criteria
* left alignment is used
* figure title below image
* figure titles are in numerical order
* appropriate alt text (descriptive, no phrases such as ‘image of’)
 
CATEGORY 7: Font
minimum font size Arial point 12
no use of italics unless direct source
no use of bold and underlining together
no use of all caps in the title
all headings should be bold or semi-bold
 
CATEGORY 8: Titles, capitalisation, numbers and dates
* titles are under 60 characters
* titles are unique
* consistent titling style
* titles are descriptive and give context
* colons used to break up long titles
* does not contain slashes or dashes
* no acronyms used
* only first word and proper nouns capitalised
* only first words and proper nouns are capitalised
* capitalise and do not capitalise list has been reviewed
* titles of species genus or higher in capitals
* all species written in italics
* titles of acts, books, reports should be in capitals
* all titles are set in single quotation marks
* numbers between zero and nine are written alphabetically (excluding exceptions)
* ordinal numbers between first and ninth are written alphabetically
* comma used in long numerals
* numbers written alphabetically at the beginning of senetences
* numerals used with measurements
* common fractions written alphabetically
* ‘to’ or ‘and’ used in number ranges
* million written alphabetically
* currency symbols used for money
* dates in format 1 October 2023
* ‘to’ used in date ranges not a hyphen unless in title
* no apostrophes used when referring to decades
 
CATEGORY 9: TABLES
* rows should not break across pages
* header row should be repeated
* font size and spacing style guide rules should apply to text in tables
 
CATEGORY 10: REFERENCING AND FOOTNOTES
* ensure references are in a modified version of the Vancouver referencing style, with the in-text references shown as [1] or [1,2,3], in square brackets.
* first reference to a source should show all publication information
* citations in bibliography should have hanging indent
* bibliography citations are written in full
* bibliography alphabetically ordered by author surname
* footnote number in superscript
* footnote number is place at end of sentence
* footnote number placed after punctuation
* all footnote numbers in the text are accompanied with a source or explanation at the bottom of the page
* if end notes or citation list is used then all footnotes should be here
* multiple end notes are used if there is more than one source
* no ‘bare’ URL links in main text or footnotes – I.E., no links starting with http or www.
* no links that take up a full bullet point
* direct quotations are in “double quotation marks”
* quotations are not in italics
* no punctuation outside of quotation marks
* square brackets or footnotes used to explain context of a sentence

STEP 2 -
Amend the document to correct each error you identify in line with the rules above. Correct any text that does not comply with these rules. Highlight every single change you make in yellow so that I can easily review your work. The output should be a docx file.

YOUR RULES: You must follow these rules without deviation.
Go through the entire document section by section. For every correction you make, you absolutely MUST apply a yellow highlight to the new text. If you are unable to make a change directly, insert a comment explaining what needs to be fixed according to the rules above."""

# ==========================================
# HELPER FUNCTIONS
# ==========================================
def chunk_document_by_words(paragraphs: list, max_words: int = 2500) -> list:
    """Groups paragraph blocks cleanly based on an absolute word-count metric."""
    chunks = []
    current_chunk = []
    current_word_count = 0
    
    for p_text in paragraphs:
        if not p_text.strip():
            continue
        p_word_count = len(p_text.split())
        
        if current_word_count + p_word_count > max_words:
            if current_chunk:
                chunks.append("\n\n".join(current_chunk))
            current_chunk = [p_text]
            current_word_count = p_word_count
        else:
            current_chunk.append(p_text)
            current_word_count += p_word_count
            
    if current_chunk:
        chunks.append("\n\n".join(current_chunk))
    return chunks

def generate_highlighted_docx(original_text: str, corrected_text: str) -> io.BytesIO:
    doc = Document()
    orig_words = original_text.split()
    corr_words = corrected_text.split()
    
    p = doc.add_paragraph()
    matcher = difflib.SequenceMatcher(None, orig_words, corr_words)
    
    for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
        if opcode == 'equal':
            p.add_run(" ".join(corr_words[j1:j2]) + " ")
        elif opcode in ('replace', 'insert'):
            run = p.add_run(" ".join(corr_words[j1:j2]) + " ")
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif opcode == 'delete':
            run = p.add_run("[...] ")
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ==========================================
# STREAMLIT UI LAYOUT
# ==========================================
st.set_page_config(page_title="ESS Document Proofreader", page_icon="⚖️", layout="wide")

st.title("Environmental Standards Scotland (ESS)")
st.subheader("Enterprise Document Compliance & Style Guide Proofreader")
st.write("Upload a corporate Microsoft Word document (`.docx`) to automatically evaluate adherence to the official ESS accessibility, tone, and formatting rules.")

uploaded_file = st.file_uploader("Select compliance document", type=["docx"])

if uploaded_file is not None:
    st.success("Document uploaded successfully.")
    
    if st.button("Execute Compliance Proofreading", type="primary"):
        github_token = os.getenv("GITHUB_TOKEN")
        
        if not github_token:
            st.error("Infrastructure Configuration Error: Missing GITHUB_TOKEN environment variable.")
            st.stop()
            
        with st.spinner("Parsing document structure and processing chunks..."):
            try:
                doc = Document(uploaded_file)
                raw_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                
                if not raw_paragraphs:
                    st.warning("The uploaded document contains no readable paragraph blocks.")
                    st.stop()
                
                # Using the revised pure word-counting strategy
                chunks = chunk_document_by_words(raw_paragraphs, max_words=2500)
                
                client = OpenAI(
                    base_url="https://models.inference.ai.azure.com",
                    api_key=github_token
                )
                
                compiled_reports = []
                compiled_corrections = []
                
                progress_bar = st.progress(0)
                for index, chunk in enumerate(chunks):
                    user_instruction = (
                        f"Analyze this document segment:\n\n{chunk}\n\n"
                        "Provide your output precisely formatted as follows:\n"
                        "---DIAGNOSTIC_REPORT_START---\n(Your Step 1 Diagnostic Report items)\n---DIAGNOSTIC_REPORT_END---\n"
                        "---CORRECTED_TEXT_START---\n(The complete corrected text of this segment)\n---CORRECTED_TEXT_END---"
                    )
                    
                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": SYSTEM_PROMPT},
                            {"role": "user", "content": user_instruction}
                        ],
                        temperature=0.0
                    )
                    
                    response_text = response.choices[0].message.content
                    
                    try:
                        report_part = response_text.split("---DIAGNOSTIC_REPORT_START---")[1].split("---DIAGNOSTIC_REPORT_END---")[0].strip()
                        text_part = response_text.split("---CORRECTED_TEXT_START---")[1].split("---CORRECTED_TEXT_END---")[0].strip()
                    except IndexError:
                        report_part = response_text
                        text_part = chunk
                    
                    compiled_reports.append(report_part)
                    compiled_corrections.append(text_part)
                    
                    progress_bar.progress((index + 1) / len(chunks))
                
                full_original = "\n\n".join(raw_paragraphs)
                full_report = "\n\n".join(compiled_reports)
                full_corrected = "\n\n".join(compiled_corrections)
                
                st.write("---")
                st.header("Step 1: Diagnostic Compliance Report")
                st.markdown(full_report)
                
                st.write("---")
                st.header("Step 2: Corrected Document Generation")
                
                final_docx_bytes = generate_highlighted_docx(full_original, full_corrected)
                
                st.download_button(
                    label="Download Corrected Document (.docx)",
                    data=final_docx_bytes,
                    file_name="ESS_Compliance_Corrected.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"An infrastructure or processing fault occurred: {str(e)}")

# ==========================================
# ENTERPRISE BRANDING FOOTER
# ==========================================
st.markdown(
    """
    <style>
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: #f0f2f6;
        color: #31333F;
        text-align: center;
        padding: 10px;
        font-size: 14px;
        font-weight: bold;
        border-top: 1px solid #e0e0e0;
        z-index: 999;
    }
    </style>
    <div class="footer">
        🛡️ Secure Enterprise Pipeline | Powered by Microsoft Azure/Copilot
    </div>
    """,
    unsafe_allow_html=True
)
